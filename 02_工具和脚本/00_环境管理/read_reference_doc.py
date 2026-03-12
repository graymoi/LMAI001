import os
from docx import Document

# 使用绝对路径
file_path = r"F:\AIlm\000项目管理\参考 政府采购项目采购需求（十五五规划）.docx"
output_path = r"F:\AIlm\000项目管理\参考_政府采购项目采购需求_关键信息拆分.md"

# 读取政府采购项目采购需求文件
if os.path.exists(file_path):
    try:
        doc = Document(file_path)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        
        content = '\n'.join(fullText)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(f"✅ 成功读取并保存：{output_path}")
        print(f"   文档段落数：{len(doc.paragraphs)}")
        print(f"   文档字符数：{len(content)}")
    except Exception as e:
        print(f"❌ 读取出错：{e}")
else:
    print(f"❌ 文件不存在：{file_path}")
